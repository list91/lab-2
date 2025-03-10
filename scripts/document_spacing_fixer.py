#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class DocumentSpacingFixer:
    """
    Класс для исправления отступов и интервалов в документе Word
    после его создания основным форматером.
    """
    
    def __init__(self, document_path):
        """Инициализация с путем к документу"""
        self.document_path = document_path
        self.document = Document(document_path)
    
    def fix_paragraph_spacing(self):
        """Исправление отступов между параграфами"""
        for paragraph in self.document.paragraphs:
            style_name = paragraph.style.name
            
            # Настройка отступов в зависимости от стиля
            if style_name == 'ВКР Глава-Раздел':
                # Заголовки глав
                paragraph.paragraph_format.space_before = Pt(24)
                paragraph.paragraph_format.space_after = Pt(18)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.first_line_indent = Mm(0)  # Без отступа первой строки
            
            elif style_name == 'ВКР Параграф':
                # Заголовки параграфов
                paragraph.paragraph_format.space_before = Pt(18)
                paragraph.paragraph_format.space_after = Pt(12)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                paragraph.paragraph_format.first_line_indent = Mm(0)  # Без отступа первой строки
            
            elif style_name == 'ВКР Пункт':
                # Заголовки пунктов
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                paragraph.paragraph_format.first_line_indent = Mm(0)  # Без отступа первой строки
            
            elif style_name == 'ВКР Обычный':
                # Обычный текст
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Mm(12.5)  # Отступ первой строки 1.25 см
                
                # Проверка на маркированный список
                if paragraph.text.startswith('•'):
                    paragraph.paragraph_format.first_line_indent = Mm(0)  # Без отступа первой строки
                    paragraph.paragraph_format.left_indent = Mm(12.5)  # Отступ слева для списка
    
    def fix_font_properties(self):
        """Исправление свойств шрифта для всех элементов"""
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                # Установка шрифта Times New Roman для всего текста
                run.font.name = 'Times New Roman'
                
                # Установка размера шрифта в зависимости от стиля параграфа
                style_name = paragraph.style.name
                if style_name == 'ВКР Глава-Раздел':
                    run.font.size = Pt(20)
                    run.font.bold = True
                elif style_name == 'ВКР Параграф':
                    run.font.size = Pt(18)
                    run.font.bold = True
                elif style_name == 'ВКР Пункт':
                    run.font.size = Pt(16)
                    run.font.bold = True
                else:
                    run.font.size = Pt(16)
                    run.font.bold = False
    
    def fix_line_spacing(self):
        """Исправление межстрочных интервалов"""
        for paragraph in self.document.paragraphs:
            # Установка межстрочного интервала 1.5 для всех параграфов
            paragraph.paragraph_format.line_spacing = 1.5
    
    def fix_page_margins(self):
        """Исправление полей страницы"""
        for section in self.document.sections:
            section.left_margin = Mm(30)    # Левое поле 3 см
            section.right_margin = Mm(15)   # Правое поле 1.5 см
            section.top_margin = Mm(20)     # Верхнее поле 2 см
            section.bottom_margin = Mm(20)  # Нижнее поле 2 см
    
    def fix_document_spacing(self):
        """Применение всех исправлений к документу"""
        self.fix_paragraph_spacing()
        self.fix_font_properties()
        self.fix_line_spacing()
        self.fix_page_margins()
        
        # Сохранение исправленного документа
        self.document.save(self.document_path)
        print(f"Отступы и интервалы в документе {self.document_path} исправлены")

def main():
    """Основная функция для запуска исправления отступов"""
    document_path = '/home/user/study/diplom/diploma.docx'
    
    # Проверка существования файла
    if not os.path.exists(document_path):
        print(f"Ошибка: файл {document_path} не найден")
        return
    
    # Создание объекта для исправления отступов
    fixer = DocumentSpacingFixer(document_path)
    fixer.fix_document_spacing()

if __name__ == '__main__':
    main()
