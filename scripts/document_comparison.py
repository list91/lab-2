import docx
from docx.shared import Pt, RGBColor
import os

def compare_documents(template_path, generated_path):
    """Детальное сравнение документов"""
    template_doc = docx.Document(template_path)
    generated_doc = docx.Document(generated_path)

    print("🔍 Сравнение документов:\n")

    # Сравнение стилей
    print("📋 Стили документа:")
    template_styles = {style.name for style in template_doc.styles}
    generated_styles = {style.name for style in generated_doc.styles}
    
    missing_styles = template_styles - generated_styles
    if missing_styles:
        print(f"❌ Отсутствующие стили: {missing_styles}")

    # Сравнение форматирования
    print("\n🖋️ Форматирование текста:")
    def analyze_paragraph_formatting(paragraphs, doc_type):
        formatting_issues = []
        for i, paragraph in enumerate(paragraphs):
            if paragraph.runs:
                run = paragraph.runs[0]
                if run.font.name != 'Times New Roman':
                    formatting_issues.append(f"{doc_type} абзац {i+1}: Шрифт {run.font.name}")
                if run.font.size and run.font.size.pt != 16:
                    formatting_issues.append(f"{doc_type} абзац {i+1}: Размер шрифта {run.font.size.pt}")
        return formatting_issues

    template_formatting_issues = analyze_paragraph_formatting(template_doc.paragraphs, "Шаблон")
    generated_formatting_issues = analyze_paragraph_formatting(generated_doc.paragraphs, "Сгенерированный")

    if template_formatting_issues:
        print("❌ Проблемы форматирования в шаблоне:")
        for issue in template_formatting_issues[:5]:
            print(issue)

    if generated_formatting_issues:
        print("\n❌ Проблемы форматирования в сгенерированном документе:")
        for issue in generated_formatting_issues[:5]:
            print(issue)

    # Сравнение цветов текста
    print("\n🎨 Цвет текста:")
    def analyze_text_color(paragraphs, doc_type):
        color_issues = []
        for i, paragraph in enumerate(paragraphs):
            for run in paragraph.runs:
                if run.font.color and run.font.color.rgb != RGBColor(0, 0, 0):
                    color_issues.append(f"{doc_type} абзац {i+1}: Цвет {run.font.color.rgb}")
        return color_issues

    template_color_issues = analyze_text_color(template_doc.paragraphs, "Шаблон")
    generated_color_issues = analyze_text_color(generated_doc.paragraphs, "Сгенерированный")

    if template_color_issues:
        print("❌ Цветовые особенности в шаблоне:")
        for issue in template_color_issues[:5]:
            print(issue)

    if generated_color_issues:
        print("\n❌ Цветовые особенности в сгенерированном документе:")
        for issue in generated_color_issues[:5]:
            print(issue)

    # Сравнение количества параграфов и текста
    print(f"\n📊 Метрики документов:")
    print(f"Шаблон: {len(template_doc.paragraphs)} параграфов")
    print(f"Сгенерированный: {len(generated_doc.paragraphs)} параграфов")

    template_text = ' '.join(p.text for p in template_doc.paragraphs if p.text.strip())
    generated_text = ' '.join(p.text for p in generated_doc.paragraphs if p.text.strip())

    print(f"Шаблон: {len(template_text)} символов")
    print(f"Сгенерированный: {len(generated_text)} символов")

def main():
    template_path = '/home/user/Downloads/vkr-2024.docx'
    generated_path = '/home/user/study/diplom/diploma.docx'
    compare_documents(template_path, generated_path)

if __name__ == '__main__':
    main()
