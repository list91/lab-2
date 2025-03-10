import os
import re
from docx import Document
from typing import List, Dict

class DiplomaStyleChecker:
    def __init__(self, directory: str):
        self.directory = directory
        self.errors = []

    def check_font(self, file_path: str) -> List[str]:
        """Проверка шрифта и его размера"""
        doc = Document(file_path)
        font_errors = []
        
        for paragraph in doc.paragraphs:
            if paragraph.style and paragraph.runs:
                run = paragraph.runs[0]
                if run.font.name != 'Times New Roman':
                    font_errors.append(f"Неверный шрифт в параграфе: {paragraph.text[:50]}...")
                if run.font.size != 14:
                    font_errors.append(f"Неверный размер шрифта в параграфе: {paragraph.text[:50]}...")
        
        return font_errors

    def check_line_spacing(self, file_path: str) -> List[str]:
        """Проверка межстрочного интервала"""
        doc = Document(file_path)
        spacing_errors = []
        
        for paragraph in doc.paragraphs:
            if paragraph.paragraph_format.line_spacing != 1.5:
                spacing_errors.append(f"Неверный межстрочный интервал в параграфе: {paragraph.text[:50]}...")
        
        return spacing_errors

    def check_margins(self, file_path: str) -> List[str]:
        """Проверка полей документа"""
        doc = Document(file_path)
        margin_errors = []
        
        sections = doc.sections
        for section in sections:
            if section.left_margin.inches != 1.18:  # 30 мм
                margin_errors.append("Левое поле не соответствует 30 мм")
            if section.right_margin.inches != 0.39:  # 10 мм
                margin_errors.append("Правое поле не соответствует 10 мм")
            if section.top_margin.inches != 0.79:  # 20 мм
                margin_errors.append("Верхнее поле не соответствует 20 мм")
            if section.bottom_margin.inches != 0.79:  # 20 мм
                margin_errors.append("Нижнее поле не соответствует 20 мм")
        
        return margin_errors

    def check_alignment(self, file_path: str) -> List[str]:
        """Проверка выравнивания по ширине"""
        doc = Document(file_path)
        alignment_errors = []
        
        for paragraph in doc.paragraphs:
            if paragraph.alignment != 3:  # WD_ALIGN_PARAGRAPH.JUSTIFY
                alignment_errors.append(f"Неверное выравнивание в параграфе: {paragraph.text[:50]}...")
        
        return alignment_errors

    def scan_documents(self) -> Dict[str, List[str]]:
        """Сканирование всех документов в директории"""
        results = {}
        
        for root, _, files in os.walk(self.directory):
            for file in files:
                if file.endswith('.docx'):
                    full_path = os.path.join(root, file)
                    file_errors = []
                    
                    file_errors.extend(self.check_font(full_path))
                    file_errors.extend(self.check_line_spacing(full_path))
                    file_errors.extend(self.check_margins(full_path))
                    file_errors.extend(self.check_alignment(full_path))
                    
                    if file_errors:
                        results[full_path] = file_errors
        
        return results

def main():
    checker = DiplomaStyleChecker('/home/user/study/diplom')
    errors = checker.scan_documents()
    
    if errors:
        print("Обнаружены ошибки в оформлении:")
        for file, file_errors in errors.items():
            print(f"\nФайл: {file}")
            for error in file_errors:
                print(f"  - {error}")
    else:
        print("Документ соответствует техническим требованиям.")

if __name__ == '__main__':
    main()
