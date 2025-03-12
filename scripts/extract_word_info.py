#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import json
import argparse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def extract_text(doc):
    """Извлекает весь текст из документа Word."""
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def extract_chapter_structure(doc):
    """Извлекает структуру глав из документа Word."""
    structure = []
    current_chapter = None
    current_subchapter = None
    
    # Предполагаем, что заголовки глав имеют стиль 'Heading 1', а подглавы - 'Heading 2'
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading 1') or para.style.name == 'Title':
            current_chapter = {
                "title": para.text,
                "level": 1,
                "subchapters": []
            }
            structure.append(current_chapter)
            current_subchapter = None
        elif para.style.name.startswith('Heading 2'):
            if current_chapter is not None:
                current_subchapter = {
                    "title": para.text,
                    "level": 2,
                    "subchapters": []
                }
                current_chapter["subchapters"].append(current_subchapter)
        elif para.style.name.startswith('Heading 3'):
            if current_subchapter is not None:
                subsubchapter = {
                    "title": para.text,
                    "level": 3
                }
                current_subchapter["subchapters"].append(subsubchapter)
    
    return structure

def extract_styles(doc):
    """Извлекает информацию о стилях из документа Word."""
    styles_info = {
        "paragraph_styles": {},
        "character_styles": {},
        "table_styles": {},
        "document_defaults": {
            "font": doc.styles['Normal'].font.name if hasattr(doc.styles['Normal'], 'font') and hasattr(doc.styles['Normal'].font, 'name') else "Unknown",
            "font_size": str(doc.styles['Normal'].font.size.pt) + " pt" if hasattr(doc.styles['Normal'], 'font') and hasattr(doc.styles['Normal'].font, 'size') else "Unknown"
        }
    }
    
    # Сбор информации о стилях абзацев
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_info = {
                "name": style.name,
                "font": style.font.name if hasattr(style, 'font') and hasattr(style.font, 'name') else "Unknown",
                "font_size": str(style.font.size.pt) + " pt" if hasattr(style, 'font') and hasattr(style.font, 'size') and style.font.size else "Unknown",
                "bold": style.font.bold if hasattr(style, 'font') and hasattr(style.font, 'bold') else False,
                "italic": style.font.italic if hasattr(style, 'font') and hasattr(style.font, 'italic') else False,
                "underline": style.font.underline if hasattr(style, 'font') and hasattr(style.font, 'underline') else False
            }
            styles_info["paragraph_styles"][style.name] = style_info
    
    # Дополнительно собираем информацию о форматировании из параграфов
    paragraph_formats = {}
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():  # Только для непустых параграфов
            para_format = para.paragraph_format
            alignment = "Unknown"
            if hasattr(para_format, 'alignment') and para_format.alignment is not None:
                if para_format.alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                    alignment = "По левому краю"
                elif para_format.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                    alignment = "По центру"
                elif para_format.alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                    alignment = "По правому краю"
                elif para_format.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                    alignment = "По ширине"
            
            left_indent = str(para_format.left_indent.inches) + " inches" if hasattr(para_format, 'left_indent') and para_format.left_indent else "0"
            right_indent = str(para_format.right_indent.inches) + " inches" if hasattr(para_format, 'right_indent') and para_format.right_indent else "0"
            first_line_indent = str(para_format.first_line_indent.inches) + " inches" if hasattr(para_format, 'first_line_indent') and para_format.first_line_indent else "0"
            
            paragraph_formats[f"paragraph_{i+1}"] = {
                "text_preview": para.text[:50] + "..." if len(para.text) > 50 else para.text,
                "style": para.style.name,
                "alignment": alignment,
                "left_indent": left_indent,
                "right_indent": right_indent,
                "first_line_indent": first_line_indent,
                "line_spacing": str(para_format.line_spacing) if hasattr(para_format, 'line_spacing') and para_format.line_spacing else "Unknown"
            }
    
    styles_info["paragraph_formats"] = paragraph_formats
    
    return styles_info

def process_word_document(input_file, output_dir):
    """Обрабатывает документ Word и сохраняет извлеченную информацию."""
    # Проверяем существование выходной директории
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Загружаем документ
    doc = Document(input_file)
    
    # Извлекаем текст
    text = extract_text(doc)
    text_file = os.path.join(output_dir, "extracted_text.txt")
    with open(text_file, 'w', encoding='utf-8') as f:
        f.write(text)
    
    # Извлекаем структуру глав
    chapters = extract_chapter_structure(doc)
    chapters_file = os.path.join(output_dir, "chapter_structure.json")
    with open(chapters_file, 'w', encoding='utf-8') as f:
        json.dump(chapters, f, ensure_ascii=False, indent=4)
    
    # Извлекаем стили
    styles = extract_styles(doc)
    styles_file = os.path.join(output_dir, "document_styles.json")
    with open(styles_file, 'w', encoding='utf-8') as f:
        json.dump(styles, f, ensure_ascii=False, indent=4)
    
    return {
        "text_file": text_file,
        "chapters_file": chapters_file,
        "styles_file": styles_file
    }

def main():
    parser = argparse.ArgumentParser(description='Извлечение информации из документа Word.')
    parser.add_argument('input_file', help='Путь к документу Word')
    parser.add_argument('output_dir', help='Путь к директории для сохранения извлеченной информации')
    
    args = parser.parse_args()
    
    try:
        results = process_word_document(args.input_file, args.output_dir)
        print(f"Обработка завершена успешно!")
        print(f"Текст сохранен в: {results['text_file']}")
        print(f"Структура глав сохранена в: {results['chapters_file']}")
        print(f"Информация о стилях сохранена в: {results['styles_file']}")
    except Exception as e:
        print(f"Ошибка при обработке документа: {str(e)}")

if __name__ == "__main__":
    main()
