import os
import re
import shutil
from typing import List, Dict, Any
import markdown
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.styles.style import _ParagraphStyle, _CharacterStyle, _TableStyle
from docxtpl import DocxTemplate

class DiplomaFormatter:
    CHAPTER_TRANSLATIONS = {
        '1_introduction': '1. Введение',
        '2_theoretical_part': '2. Теоретическая часть',
        '3_practical_implementation': '3. Практическая реализация',
        '4_research_methodology': '4. Методология исследования',
        '5_research_results': '5. Результаты исследования',
        '6_practical_significance': '6. Практическая значимость',
        '7_development_prospects': '7. Перспективы развития',
        '8_appendices': '8. Приложения'
    }

    def __init__(self, chapters_dir: str, output_path: str, template_path: str):
        self.chapters_dir = chapters_dir
        self.output_path = output_path
        self.template_path = template_path
        
        # Создаем копию шаблона для работы
        shutil.copy(self.template_path, self.output_path)
        
        # Открываем документ на основе шаблона
        self.document = Document(self.output_path)
        
        # Очищаем содержимое шаблона, но сохраняем стили
        self._clear_template_content()
        
        # Настраиваем дополнительные стили, если нужно
        self._setup_additional_styles()

    def _clear_template_content(self):
        """Очистка содержимого шаблона, сохраняя стили и структуру"""
        # Удаляем все параграфы, кроме последнего
        while len(self.document.paragraphs) > 1:
            p = self.document.paragraphs[0]._p
            p.getparent().remove(p)
            
        # Если остался один параграф, очищаем его
        if len(self.document.paragraphs) == 1:
            self.document.paragraphs[0].text = ''
            
        # Удаляем все таблицы
        while len(self.document.tables) > 0:
            tbl = self.document.tables[0]._tbl
            tbl.getparent().remove(tbl)
            
        # Добавляем пустой параграф для начала документа
        self.document.add_paragraph()

    def _setup_additional_styles(self):
        """Настройка дополнительных стилей, если они не определены в шаблоне"""
        # Проверяем наличие основных стилей
        required_styles = ['ВКР Обычный', 'ВКР Глава-Раздел', 'ВКР Параграф', 'ВКР Пункт']
        
        for style_name in required_styles:
            if style_name not in [s.name for s in self.document.styles]:
                # Если стиль отсутствует, создаем его на основе базовых стилей
                if style_name == 'ВКР Обычный':
                    style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(16)
                    style.paragraph_format.line_spacing = 1.5
                    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    style.paragraph_format.first_line_indent = Mm(12.5)
                    
                elif style_name == 'ВКР Глава-Раздел':
                    style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(20)
                    style.font.bold = True
                    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    style.paragraph_format.space_after = Pt(12)
                    
                elif style_name == 'ВКР Параграф':
                    style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(18)
                    style.font.bold = True
                    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    style.paragraph_format.space_after = Pt(12)
                    
                elif style_name == 'ВКР Пункт':
                    style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    style.font.name = 'Times New Roman'
                    style.font.size = Pt(16)
                    style.font.bold = True
                    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    style.paragraph_format.space_after = Pt(12)

    def _convert_markdown_to_docx(self, markdown_text: str):
        """Конвертация Markdown в docx с сохранением структуры"""
        # Обработка заголовков второго и третьего уровня
        lines = markdown_text.split('\n')
        processed_lines = []
        
        for line in lines:
            if line.startswith('## '):
                # Заголовок второго уровня (параграф)
                heading_text = line.replace('## ', '').strip()
                # Используем стиль ВКР Параграф вместо стандартного Heading 2
                if 'ВКР Параграф' in [s.name for s in self.document.styles]:
                    self.document.add_paragraph(heading_text, style='ВКР Параграф')
                else:
                    self.document.add_heading(heading_text, level=2)
            elif line.startswith('### '):
                # Заголовок третьего уровня (пункт)
                heading_text = line.replace('### ', '').strip()
                # Используем стиль ВКР Пункт вместо стандартного Heading 3
                if 'ВКР Пункт' in [s.name for s in self.document.styles]:
                    self.document.add_paragraph(heading_text, style='ВКР Пункт')
                else:
                    self.document.add_heading(heading_text, level=3)
            else:
                processed_lines.append(line)
        
        # Собираем обратно текст без заголовков
        markdown_text = '\n'.join(processed_lines)
        
        # Обработка списков
        markdown_text = re.sub(r'^- ', '• ', markdown_text, flags=re.MULTILINE)
        
        # Обработка кода
        markdown_text = re.sub(r'```(.*?)```', r'[Листинг кода]\n\1', markdown_text, flags=re.DOTALL)
        
        # Преобразуем в HTML
        html = markdown.markdown(markdown_text)
        
        # Разбор HTML и добавление в документ
        paragraphs = re.split(r'<p>|</p>', html)
        paragraphs = [p for p in paragraphs if p.strip()]
        
        for paragraph in paragraphs:
            # Удаление HTML-тегов
            clean_text = re.sub(r'<[^>]+>', '', paragraph).strip()
            
            if clean_text:
                # Используем стиль ВКР Обычный вместо стандартного Normal
                if 'ВКР Обычный' in [s.name for s in self.document.styles]:
                    para = self.document.add_paragraph(clean_text, style='ВКР Обычный')
                else:
                    para = self.document.add_paragraph(clean_text, style='Normal')
                
                # Явно устанавливаем шрифт для каждого фрагмента текста
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)

    def _process_chapter(self, chapter_path: str):
        """Обработка главы"""
        with open(chapter_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Определение имени главы
        path_parts = chapter_path.split('/')
        
        # Проверяем структуру пути
        if path_parts[-1] == 'content.md':
            # Определяем, находится ли файл непосредственно в директории главы или в подкаталоге
            if 'content.md' in path_parts[-1] and path_parts[-2] in ['6_practical_significance', '7_development_prospects', '8_appendices']:
                chapter_name = path_parts[-2]  # Берем имя директории главы
            else:
                # Для подразделов (например, 1.1, 2.1, и т.д.)
                chapter_name = path_parts[-3]  # Берем имя директории главы
        
        # Перевод заголовка главы
        translated_name = self.CHAPTER_TRANSLATIONS.get(chapter_name, chapter_name)
        
        # Добавление заголовка главы с использованием специального стиля
        if 'ВКР Глава-Раздел' in [s.name for s in self.document.styles]:
            self.document.add_paragraph(translated_name, style='ВКР Глава-Раздел')
        else:
            self.document.add_heading(translated_name, level=1)
        
        # Конвертация контента
        self._convert_markdown_to_docx(content)
        
        # Разрыв страницы после главы
        self.document.add_page_break()

    def compile_diploma(self):
        """Компиляция всего диплома"""
        # Определяем порядок глав
        chapter_order = [
            '1_introduction',
            '2_theoretical_part',
            '3_practical_implementation',
            '4_research_methodology',
            '5_research_results',
            '6_practical_significance',
            '7_development_prospects',
            '8_appendices'
        ]
        
        # Находим все файлы content.md
        all_content_files = [
            os.path.join(root, file)
            for root, _, files in os.walk(self.chapters_dir)
            for file in files if file == 'content.md'
        ]
        
        # Обрабатываем главы в нужном порядке
        for chapter_name in chapter_order:
            # Находим все файлы content.md для текущей главы
            chapter_files = [path for path in all_content_files if f'/{chapter_name}/' in path]
            
            # Сортируем подразделы, если они есть
            chapter_files.sort()
            
            # Обработка файлов главы
            for chapter_path in chapter_files:
                self._process_chapter(chapter_path)

        # Сохранение документа
        self.document.save(self.output_path)
        print(f"Диплом сохранен в {self.output_path}")

def main():
    diploma_dir = '/home/user/study/diplom/chapters'
    output_path = '/home/user/study/diplom/diploma.docx'
    template_path = '/home/user/Downloads/vkr-2024.docx'
    
    formatter = DiplomaFormatter(diploma_dir, output_path, template_path)
    formatter.compile_diploma()

if __name__ == '__main__':
    main()
