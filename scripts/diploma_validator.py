import os
import re
from typing import Dict, List, Any
import docx
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING

class DiplomaValidator:
    EXPECTED_CHAPTERS = [
        '1. Введение',
        '2. Теоретическая часть',
        '3. Практическая реализация',
        '4. Методология исследования',
        '5. Результаты исследования',
        '6. Практическая значимость',
        '7. Перспективы развития',
        '8. Приложения'
    ]

    def __init__(self, document_path: str):
        self.document = docx.Document(document_path)
        self.validation_results = {
            'структурные_требования': [],
            'технические_требования': [],
            'стилистические_замечания': [],
            'метрики_документа': {}
        }

    def check_document_structure(self):
        """Проверка структуры документа"""
        # Получаем все заголовки первого уровня
        headings = []
        for p in self.document.paragraphs:
            if p.style and p.style.name == 'Heading 1':
                headings.append(p.text)
            elif p.style and p.style.name.startswith('Heading') and p.text.strip().startswith(tuple('12345678')):
                headings.append(p.text)

        # Проверка наличия глав
        found_chapters = []
        for expected_chapter in self.EXPECTED_CHAPTERS:
            chapter_num = expected_chapter.split('.')[0]
            found = False
            for heading in headings:
                if heading.startswith(chapter_num + '.') or heading.startswith(chapter_num + ' '):
                    found = True
                    found_chapters.append(expected_chapter)
                    break
            
            if not found:
                self.validation_results['структурные_требования'].append(
                    f'❌ Отсутствует глава: {expected_chapter}'
                )

        # Проверка порядка глав
        if found_chapters != [ch for ch in self.EXPECTED_CHAPTERS if ch in found_chapters]:
            self.validation_results['структурные_требования'].append(
                '❌ Нарушен порядок глав в документе'
            )

    def check_technical_requirements(self):
        """Проверка технических требований ГОСТ"""
        sections = self.document.sections
        for section in sections:
            # Проверка полей
            checks = [
                (section.left_margin.inches, 1.18, 'Левое поле'),
                (section.right_margin.inches, 0.59, 'Правое поле'),
                (section.top_margin.inches, 0.79, 'Верхнее поле'),
                (section.bottom_margin.inches, 0.79, 'Нижнее поле')
            ]

            for current, expected, name in checks:
                if abs(current * 25.4 - expected * 25.4) > 1:  # Погрешность 1 мм
                    self.validation_results['технические_требования'].append(
                        f'❌ {name} не соответствует требованиям (текущее: {current * 25.4:.2f} мм, требуется: {expected * 25.4:.2f} мм)'
                    )

    def check_typography(self):
        """Проверка типографских требований"""
        font_errors = 0
        font_error_details = []
        
        # Проверка стилей абзацев
        for i, paragraph in enumerate(self.document.paragraphs):
            # Пропускаем пустые абзацы
            if not paragraph.text.strip():
                continue
                
            # Пропускаем заголовки
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                continue
                
            # Проверка шрифта в каждом фрагменте текста
            for run in paragraph.runs:
                # Проверка шрифта
                if run.font.name and run.font.name != 'Times New Roman':
                    font_errors += 1
                    if len(font_error_details) < 5:  # Сохраняем только первые 5 ошибок для примера
                        font_error_details.append(f'Абзац {i+1}: Шрифт {run.font.name} вместо Times New Roman')
                
                # Проверка размера шрифта
                if run.font.size and run.font.size.pt != 16:
                    font_errors += 1
                    if len(font_error_details) < 5:
                        font_error_details.append(f'Абзац {i+1}: Размер шрифта {run.font.size.pt} вместо 16')

        # Добавляем информацию об ошибках
        if font_errors > 0:
            self.validation_results['технические_требования'].append(
                f'❌ Обнаружено {font_errors} нарушений шрифта и размера'
            )
            
            # Добавляем примеры ошибок
            if font_error_details:
                self.validation_results['технические_требования'].append(
                    f'ℹ️ Примеры ошибок:\n' + '\n'.join(font_error_details) + 
                    (f'\n... и еще {font_errors - len(font_error_details)} ошибок' if font_errors > len(font_error_details) else '')
                )

    def calculate_document_metrics(self):
        """Расчет метрик документа"""
        paragraphs = [p for p in self.document.paragraphs if p.text.strip()]
        words = sum(len(p.text.split()) for p in paragraphs)
        characters = sum(len(p.text) for p in paragraphs)

        self.validation_results['метрики_документа'] = {
            'количество_параграфов': len(paragraphs),
            'количество_слов': words,
            'количество_символов': characters,
            'приблизительное_количество_страниц': len(paragraphs) // 10
        }

    def check_formatting_consistency(self):
        """Проверка согласованности форматирования"""
        paragraph_styles = {}
        for paragraph in self.document.paragraphs:
            if paragraph.style:
                paragraph_styles[paragraph.style.name] = paragraph_styles.get(paragraph.style.name, 0) + 1

        if len(paragraph_styles) > 5:
            self.validation_results['стилистические_замечания'].append(
                f'⚠️ Слишком много различных стилей: {list(paragraph_styles.keys())}'
            )

    def validate(self):
        """Полная валидация документа"""
        self.check_document_structure()
        self.check_technical_requirements()
        self.check_typography()
        self.calculate_document_metrics()
        self.check_formatting_consistency()
        return self.validation_results

def main():
    document_path = '/home/user/study/diplom/diploma.docx'
    validator = DiplomaValidator(document_path)
    results = validator.validate()

    print("🔍 Результаты валидации диплома:\n")
    
    print("📋 Структурные требования:")
    for req in results['структурные_требования']:
        print(req)
    
    print("\n⚙️ Технические требования:")
    for req in results['технические_требования']:
        print(req)
    
    print("\n✏️ Стилистические замечания:")
    for note in results['стилистические_замечания']:
        print(note)
    
    print("\n📊 Метрики документа:")
    for metric, value in results['метрики_документа'].items():
        print(f"{metric.replace('_', ' ').capitalize()}: {value}")

if __name__ == '__main__':
    main()
