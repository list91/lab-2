#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import docx
import json

class DocumentTextExtractor:
    """
    Класс для извлечения текстового содержимого из документа Word
    с сохранением структуры и метаданных
    """
    
    def __init__(self, document_path):
        """
        Инициализация экстрактора
        
        :param document_path: Путь к документу Word
        """
        self.document_path = document_path
        self.document = docx.Document(document_path)
    
    def extract_full_text(self):
        """
        Извлечение полного текста документа
        
        :return: Полный текст документа
        """
        full_text = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        return "\n".join(full_text)
    
    def extract_structured_text(self):
        """
        Извлечение структурированного текста с сохранением иерархии
        
        :return: Словарь с иерархической структурой документа
        """
        document_structure = {
            "заголовок": "",
            "главы": []
        }
        
        current_chapter = None
        current_section = None
        
        for paragraph in self.document.paragraphs:
            # Пропускаем пустые параграфы
            if not paragraph.text.strip():
                continue
            
            # Определение стиля параграфа
            style_name = paragraph.style.name if paragraph.style else "Normal"
            
            # Обработка заголовков
            if style_name in ["Title", "ВКР Глава-Раздел"]:
                document_structure["заголовок"] = paragraph.text
            
            # Обработка глав
            elif style_name in ["Heading 1", "ВКР Параграф"]:
                current_chapter = {
                    "название": paragraph.text,
                    "разделы": []
                }
                document_structure["главы"].append(current_chapter)
                current_section = None
            
            # Обработка разделов
            elif style_name in ["Heading 2", "ВКР Пункт"]:
                if current_chapter:
                    current_section = {
                        "название": paragraph.text,
                        "параграфы": []
                    }
                    current_chapter["разделы"].append(current_section)
            
            # Обработка основного текста
            elif style_name in ["Normal", "ВКР Обычный"]:
                if current_section:
                    current_section["параграфы"].append(paragraph.text)
                elif current_chapter:
                    current_chapter["разделы"].append(paragraph.text)
        
        return document_structure
    
    def extract_metadata(self):
        """
        Извлечение метаданных документа
        
        :return: Словарь с метаданными
        """
        metadata = {
            "путь_файла": self.document_path,
            "количество_параграфов": len(self.document.paragraphs),
            "количество_таблиц": len(self.document.tables),
            "количество_изображений": len(self.document.inline_shapes),
            "стили": [paragraph.style.name for paragraph in self.document.paragraphs if paragraph.style]
        }
        
        return metadata
    
    def save_text_to_file(self, output_format='txt'):
        """
        Сохранение извлеченного текста в файл
        
        :param output_format: Формат вывода (txt, json)
        """
        # Базовый путь для сохранения отчетов
        reports_dir = '/home/user/study/diplom/reports/full_text'
        os.makedirs(reports_dir, exist_ok=True)
        
        if output_format == 'txt':
            # Сохранение полного текста
            output_path = os.path.join(reports_dir, 'full_text.txt')
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.extract_full_text())
            print(f"Полный текст сохранен в {output_path}")
        
        elif output_format == 'json':
            # Сохранение структурированного текста
            output_path = os.path.join(reports_dir, 'structured_text.json')
            structured_text = {
                "полный_текст": self.extract_full_text(),
                "структура": self.extract_structured_text(),
                "метаданные": self.extract_metadata()
            }
            
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(structured_text, f, ensure_ascii=False, indent=2)
            print(f"Структурированный текст сохранен в {output_path}")

def main():
    """Основная функция для запуска извлечения текста"""
    document_path = '/home/user/study/diplom/diploma.docx'
    
    # Проверка существования файла
    if not os.path.exists(document_path):
        print(f"Ошибка: файл {document_path} не найден")
        return
    
    # Создание объекта для извлечения текста
    extractor = DocumentTextExtractor(document_path)
    
    # Извлечение и сохранение текста в разных форматах
    extractor.save_text_to_file('txt')
    extractor.save_text_to_file('json')

if __name__ == '__main__':
    main()
