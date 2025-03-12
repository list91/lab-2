#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Скрипт для исправления проблемы с подчеркиванием текста в документе Word.
Устанавливает русский язык для всего документа и отключает проверку орфографии.
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_russian_language(doc_path):
    """
    Устанавливает русский язык для всего документа и отключает проверку орфографии.
    
    Args:
        doc_path: Путь к документу Word
    """
    # Открываем документ
    doc = Document(doc_path)
    
    # Устанавливаем русский язык для всего документа
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            # Проверяем наличие rPr и создаем его при необходимости
            if run._element.rPr is None:
                run._element.get_or_add_rPr()
            
            # Устанавливаем русский язык (ru-RU)
            lang = OxmlElement('w:lang')
            lang.set(qn('w:val'), 'ru-RU')
            run._element.rPr.append(lang)
            
            # Отключаем проверку орфографии
            no_proof = OxmlElement('w:noProof')
            no_proof.set(qn('w:val'), '1')
            run._element.rPr.append(no_proof)
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Проверяем наличие rPr и создаем его при необходимости
                        if run._element.rPr is None:
                            run._element.get_or_add_rPr()
                        
                        # Устанавливаем русский язык (ru-RU)
                        lang = OxmlElement('w:lang')
                        lang.set(qn('w:val'), 'ru-RU')
                        run._element.rPr.append(lang)
                        
                        # Отключаем проверку орфографии
                        no_proof = OxmlElement('w:noProof')
                        no_proof.set(qn('w:val'), '1')
                        run._element.rPr.append(no_proof)
    
    # Сохраняем документ
    output_path = doc_path.replace('.docx', '_fixed.docx')
    doc.save(output_path)
    print(f"Документ успешно обработан и сохранен как: {output_path}")
    
    return output_path

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        doc_path = sys.argv[1]
    else:
        doc_path = "diploma.docx"  # По умолчанию
    
    set_russian_language(doc_path)
