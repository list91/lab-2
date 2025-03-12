#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import argparse
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import markdown
from bs4 import BeautifulSoup

def setup_document_styles(doc):
    """Настройка стилей документа в соответствии с требованиями к дипломной работе."""
    
    # Настройка стиля Normal
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(14)
    style_normal.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style_normal.paragraph_format.space_after = Pt(0)
    style_normal.paragraph_format.first_line_indent = Cm(1.25)
    style_normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Настройка стиля Title (для заголовка работы)
    if 'Title' in doc.styles:
        style_title = doc.styles['Title']
    else:
        style_title = doc.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    style_title.base_style = doc.styles['Normal']
    style_title.font.bold = True
    style_title.font.size = Pt(16)
    style_title.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    style_title.paragraph_format.space_before = Pt(0)
    style_title.paragraph_format.space_after = Pt(12)
    style_title.paragraph_format.first_line_indent = Cm(0)
    
    # Настройка стиля Heading 1 (для глав)
    style_heading1 = doc.styles['Heading 1']
    style_heading1.font.name = 'Times New Roman'
    style_heading1.font.size = Pt(16)
    style_heading1.font.bold = True
    style_heading1.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_heading1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    style_heading1.paragraph_format.space_before = Pt(12)
    style_heading1.paragraph_format.space_after = Pt(12)
    style_heading1.paragraph_format.first_line_indent = Cm(0)
    style_heading1.paragraph_format.keep_with_next = True
    
    # Настройка стиля Heading 2 (для разделов) - исправлено на 13pt как в эталоне
    style_heading2 = doc.styles['Heading 2']
    style_heading2.font.name = 'Times New Roman'
    style_heading2.font.size = Pt(13)  # Изменено с 14 на 13 согласно эталону
    style_heading2.font.bold = True
    style_heading2.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_heading2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style_heading2.paragraph_format.space_before = Pt(12)
    style_heading2.paragraph_format.space_after = Pt(6)
    style_heading2.paragraph_format.first_line_indent = Cm(0)
    style_heading2.paragraph_format.keep_with_next = True
    
    # Настройка стиля Heading 3 (для подразделов) - убран курсив согласно эталону
    style_heading3 = doc.styles['Heading 3']
    style_heading3.font.name = 'Times New Roman'
    style_heading3.font.size = Pt(14)
    style_heading3.font.bold = True
    style_heading3.font.italic = False  # Убран курсив согласно эталону
    style_heading3.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_heading3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    style_heading3.paragraph_format.space_before = Pt(6)
    style_heading3.paragraph_format.space_after = Pt(6)
    style_heading3.paragraph_format.first_line_indent = Cm(0)
    style_heading3.paragraph_format.keep_with_next = True
    
    # Добавление специальных стилей из эталонного документа
    
    # Стиль для списков
    if 'List Bullet' in doc.styles:
        style_list = doc.styles['List Bullet']
    else:
        style_list = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    style_list.base_style = doc.styles['Normal']
    style_list.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_list.paragraph_format.left_indent = Cm(1.25)
    style_list.paragraph_format.first_line_indent = Cm(0)
    
    # Стиль для кода
    if 'Code' in doc.styles:
        style_code = doc.styles['Code']
    else:
        style_code = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    style_code.base_style = doc.styles['Normal']
    style_code.font.name = 'Courier New'
    style_code.font.size = Pt(12)
    style_code.font.color.rgb = RGBColor(0, 0, 128)  # Темно-синий цвет для кода
    style_code.paragraph_format.left_indent = Cm(1.25)
    style_code.paragraph_format.first_line_indent = Cm(0)
    style_code.paragraph_format.space_before = Pt(6)
    style_code.paragraph_format.space_after = Pt(6)
    
    # Стиль для сносок
    if 'footnote text' not in doc.styles:
        style_footnote = doc.styles.add_style('footnote text', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_footnote = doc.styles['footnote text']
    style_footnote.font.name = 'Times New Roman'
    style_footnote.font.size = Pt(10)
    style_footnote.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    
    # Специальные стили ВКР из эталонного документа
    
    # ВКР Содержимое таблицы
    if 'ВКР Содержимое таблицы' not in doc.styles:
        style_table_content = doc.styles.add_style('ВКР Содержимое таблицы', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_table_content = doc.styles['ВКР Содержимое таблицы']
    style_table_content.base_style = doc.styles['Normal']
    style_table_content.font.name = 'Times New Roman'
    style_table_content.font.size = Pt(12)
    style_table_content.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_table_content.paragraph_format.first_line_indent = Cm(0)
    
    # ВКР Название таблицы
    if 'ВКР Название таблицы' not in doc.styles:
        style_table_title = doc.styles.add_style('ВКР Название таблицы', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_table_title = doc.styles['ВКР Название таблицы']
    style_table_title.base_style = doc.styles['Normal']
    style_table_title.font.name = 'Times New Roman'
    style_table_title.font.size = Pt(12)
    style_table_title.font.bold = True
    style_table_title.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_table_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    style_table_title.paragraph_format.first_line_indent = Cm(0)
    
    # ВКР Рисунок
    if 'ВКР Рисунок' not in doc.styles:
        style_figure = doc.styles.add_style('ВКР Рисунок', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_figure = doc.styles['ВКР Рисунок']
    style_figure.base_style = doc.styles['Normal']
    style_figure.font.name = 'Times New Roman'
    style_figure.font.size = Pt(12)
    style_figure.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_figure.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    style_figure.paragraph_format.first_line_indent = Cm(0)
    
    # ВКР Литература
    if 'ВКР Литература' not in doc.styles:
        style_literature = doc.styles.add_style('ВКР Литература', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_literature = doc.styles['ВКР Литература']
    style_literature.base_style = doc.styles['Normal']
    style_literature.font.name = 'Times New Roman'
    style_literature.font.size = Pt(14)
    style_literature.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    style_literature.paragraph_format.first_line_indent = Cm(-0.63)
    style_literature.paragraph_format.left_indent = Cm(1.25)
    
    # Стили для оглавления
    if 'toc 1' not in doc.styles:
        style_toc1 = doc.styles.add_style('toc 1', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_toc1 = doc.styles['toc 1']
    style_toc1.base_style = doc.styles['Normal']
    style_toc1.font.name = 'Times New Roman'
    style_toc1.font.size = Pt(14)
    style_toc1.font.bold = True
    style_toc1.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    
    if 'toc 2' not in doc.styles:
        style_toc2 = doc.styles.add_style('toc 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        style_toc2 = doc.styles['toc 2']
    style_toc2.base_style = doc.styles['Normal']
    style_toc2.font.name = 'Times New Roman'
    style_toc2.font.size = Pt(14)
    style_toc2.font.color.rgb = RGBColor(0, 0, 0)  # Черный цвет
    
    return doc

def add_title_page(doc):
    """Добавление титульной страницы в документ."""
    
    # Верхняя часть титульной страницы
    p = doc.add_paragraph('МИНИСТЕРСТВО ПРОСВЕЩЕНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ ВЫСШЕГО ОБРАЗОВАНИЯ', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph('«РОССИЙСКИЙ ГОСУДАРСТВЕННЫЙ', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph('ПЕДАГОГИЧЕСКИЙ УНИВЕРСИТЕТ им. А. И. ГЕРЦЕНА»', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Пустые строки для отступа
    for _ in range(5):
        doc.add_paragraph()
    
    # Информация о направлении подготовки
    p = doc.add_paragraph('Направление подготовки/специальность', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(7)
    
    p = doc.add_paragraph('44.03.01 Педагогическое образование', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(7)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph('направленность (профиль)/специализация', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(7)
    
    p = doc.add_paragraph('«Физическое образование»', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(7)
    
    # Название работы
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph('Выпускная квалификационная работа', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph('Автоматизированная система классификации болезней томатов с использованием методов машинного обучения', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Пустые строки для отступа
    for _ in range(10):
        doc.add_paragraph()
    
    # Город и год
    p = doc.add_paragraph('Санкт-Петербург', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    p = doc.add_paragraph('2025', style='Normal')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    
    # Добавляем разрыв страницы
    doc.add_page_break()

def add_table_of_contents(doc):
    """Добавление оглавления в документ."""
    
    p = doc.add_paragraph('ОГЛАВЛЕНИЕ', style='Heading 1')
    
    # Добавляем пустой параграф с текстом-заполнителем для оглавления
    # и применяем стиль toc 1 для первого уровня
    p = doc.add_paragraph()
    p.style = 'toc 1'
    p.add_run("Здесь будет оглавление. Обновите его после открытия документа.")
    
    # Добавляем комментарий с инструкцией
    p = doc.add_paragraph()
    p.add_run("(Для обновления оглавления в Word: выделите содержимое, нажмите F9 или щелкните правой кнопкой мыши и выберите 'Обновить поле')")
    
    # Добавляем разрыв страницы
    doc.add_page_break()

def parse_markdown_file(file_path):
    """Преобразование markdown-файла в HTML для последующей обработки."""
    with open(file_path, 'r', encoding='utf-8') as file:
        md_content = file.read()
    
    # Преобразуем Markdown в HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    return html_content

def extract_chapter_info(dir_name):
    """Извлечение информации о главе из имени директории."""
    # Пример: 1_introduction -> (1, "introduction")
    match = re.match(r'(\d+)(?:\.(\d+))?_(.+)', dir_name)
    if match:
        if match.group(2):  # Если есть номер подраздела
            return int(match.group(1)), int(match.group(2)), match.group(3).replace('_', ' ')
        else:
            return int(match.group(1)), 0, match.group(3).replace('_', ' ')
    return 0, 0, dir_name.replace('_', ' ')

def add_html_content_to_doc(doc, html_content, level=1):
    """Добавление HTML-контента в документ Word с учетом стилей."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'li', 'pre', 'code', 'table']):
        if element.name == 'h1':
            p = doc.add_paragraph(element.text, style='Heading 1')
        elif element.name == 'h2':
            p = doc.add_paragraph(element.text, style='Heading 2')
        elif element.name == 'h3':
            p = doc.add_paragraph(element.text, style='Heading 3')
        elif element.name == 'p':
            p = doc.add_paragraph(element.text, style='Normal')
        elif element.name == 'ul':
            for li in element.find_all('li'):
                p = doc.add_paragraph(li.text, style='List Bullet')
                p.paragraph_format.first_line_indent = Cm(-0.63)
                p.paragraph_format.left_indent = Cm(1.25)
                run = p.add_run()
                run.add_tab()
        elif element.name == 'pre' or element.name == 'code':
            p = doc.add_paragraph(element.text, style='Code')
        elif element.name == 'table':
            # Обработка таблиц
            rows = element.find_all('tr')
            if rows:
                # Добавляем название таблицы
                caption = element.find('caption')
                if caption:
                    table_title = doc.add_paragraph(caption.text, style='ВКР Название таблицы')
                
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table_cell = table.cell(i, j)
                        table_cell.text = cell.text
                        # Применяем стиль к параграфу в ячейке
                        for paragraph in table_cell.paragraphs:
                            paragraph.style = 'ВКР Содержимое таблицы'
                        
                # Добавляем пустой параграф после таблицы
                doc.add_paragraph()
    
    return doc

def process_chapters_directory(doc, chapters_dir, level=1):
    """Рекурсивная обработка директории с главами."""
    # Получаем список директорий
    dirs = [d for d in os.listdir(chapters_dir) if os.path.isdir(os.path.join(chapters_dir, d))]
    
    # Сортируем директории по номеру главы
    dirs.sort(key=lambda x: extract_chapter_info(x)[0])
    
    for dir_name in dirs:
        chapter_num, subchapter_num, chapter_title = extract_chapter_info(dir_name)
        chapter_path = os.path.join(chapters_dir, dir_name)
        
        # Проверяем наличие файла content.md
        content_file = os.path.join(chapter_path, 'content.md')
        if os.path.exists(content_file):
            # Обрабатываем содержимое главы
            html_content = parse_markdown_file(content_file)
            doc = add_html_content_to_doc(doc, html_content, level)
        
        # Рекурсивно обрабатываем подпапки
        subdirs = [d for d in os.listdir(chapter_path) if os.path.isdir(os.path.join(chapter_path, d))]
        if subdirs:
            process_chapters_directory(doc, chapter_path, level + 1)
    
    return doc

def add_bibliography(doc):
    """Добавление списка литератур в документ."""
    p = doc.add_paragraph('СПИСОК ЛИТЕРАТУРЫ', style='Heading 1')
    
    # Здесь можно добавить список литератур из отдельного файла, если он существует
    # Или добавить стандартный шаблон списка литератур
    
    p = doc.add_paragraph('1. Иванов И.И. Название книги. - М.: Издательство, 2023. - 123 с.', style='ВКР Литература')
    
    p = doc.add_paragraph('2. Петров П.П. Название статьи // Название журнала. - 2022. - №5. - С. 10-15.', style='ВКР Литература')
    
    p = doc.add_paragraph('3. Сидоров С.С. Название диссертации: дис. ... канд. наук. - СПб., 2021. - 150 с.', style='ВКР Литература')
    
    return doc

def add_appendices(doc):
    """Добавление приложений в документ."""
    p = doc.add_paragraph('ПРИЛОЖЕНИЯ', style='Heading 1')
    
    # Здесь можно добавить приложения из отдельных файлов, если они существуют
    
    return doc

def create_word_document(chapters_dir, output_file):
    """Создание документа Word на основе содержимого директории с главами."""
    # Создаем новый документ
    doc = Document()
    
    # Настраиваем стили документа
    doc = setup_document_styles(doc)
    
    # Добавляем титульную страницу
    add_title_page(doc)
    
    # Добавляем оглавление
    add_table_of_contents(doc)
    
    # Обрабатываем директорию с главами
    doc = process_chapters_directory(doc, chapters_dir)
    
    # Добавляем список литератур
    add_bibliography(doc)
    
    # Добавляем приложения
    add_appendices(doc)
    
    # Сохраняем документ
    doc.save(output_file)
    
    return output_file

def main():
    parser = argparse.ArgumentParser(description='Создание документа Word на основе содержимого директории с главами.')
    parser.add_argument('--chapters_dir', default='chapters', help='Путь к директории с главами')
    parser.add_argument('--output_file', default='diploma.docx', help='Путь к выходному файлу')
    
    args = parser.parse_args()
    
    try:
        output_file = create_word_document(args.chapters_dir, args.output_file)
        print(f"Документ успешно создан: {output_file}")
    except Exception as e:
        print(f"Ошибка при создании документа: {str(e)}")

if __name__ == "__main__":
    main()
